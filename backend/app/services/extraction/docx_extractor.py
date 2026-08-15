from docx import Document as DocxDocument

from app.services.extraction.base_extractor import BaseExtractor, ExtractedPage

# DOCX has no reliable native page boundaries (pagination depends on the
# renderer), so we approximate a "page" every N paragraphs. This keeps the
# citation UX consistent with PDF ("page 3") without lying about precision.
PARAGRAPHS_PER_PSEUDO_PAGE = 25


class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str) -> list[ExtractedPage]:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from table cells if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text and row_text not in paragraphs:
                    paragraphs.append(row_text)

        pages: list[ExtractedPage] = []
        for i in range(0, len(paragraphs), PARAGRAPHS_PER_PSEUDO_PAGE):
            chunk = paragraphs[i : i + PARAGRAPHS_PER_PSEUDO_PAGE]
            text = self.clean_text("\n".join(chunk))
            if text:
                pages.append(ExtractedPage(text=text, page_number=(i // PARAGRAPHS_PER_PSEUDO_PAGE) + 1))
        return pages
